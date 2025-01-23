# Install necessary libraries
# pip install streamlit pandas openpyxl python-docx

import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

# Define the fixed Excel file path
file_path = "Bank Nodal Officer Email I.D.xlsx"

# Function to fetch details from the Excel file
def fetch_emails_from_excel(bank_names):
    try:
        df = pd.read_excel(file_path)  # Load the Excel file
        result = []
        
        # Iterate over the provided bank names and fetch their details
        for bank in bank_names:
            row = df[df['Bank Name'].str.lower() == bank.lower()]
            if not row.empty:
                details = {
                    'Bank Name': bank,
                    'Customer Email': row['Customer Email'].values[0] if not pd.isna(row['Customer Email'].values[0]) else '',
                    'Nodal Email': row['Nodal Email'].values[0] if not pd.isna(row['Nodal Email'].values[0]) else '',
                    'Grievance Email': row['Grievance Email'].values[0] if not pd.isna(row['Grievance Email'].values[0]) else ''
                }
            else:
                details = {'Bank Name': bank, 'Customer Email': '', 'Nodal Email': '', 'Grievance Email': ''}
            result.append(details)
        return result
    except FileNotFoundError:
        st.error("Excel file not found. Please check the file path.")
        return []

# Function to create a Word document
def create_word_file(bank_details):
    doc = Document()
    doc.add_heading("Banks Email (Nodal)", level=1)

    for details in bank_details:
        doc.add_paragraph(f"Bank Name: {details['Bank Name']}")
        doc.add_paragraph(f"Customer Email: {details['Customer Email']}")
        doc.add_paragraph(f"Nodal Email: {details['Nodal Email']}")
        doc.add_paragraph(f"Grievance Email: {details['Grievance Email']}")
        doc.add_paragraph("")  # Add a blank line after each bank's details

    # Save the Word file to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit App
def main():
    st.title("Nodal Generator")

    # Load Excel file and display it
    try:
        df = pd.read_excel(file_path)
        st.success("File loaded successfully from fixed path!")
        st.dataframe(df.head())  # Display first few rows
    except FileNotFoundError:
        st.error("Excel file not found. Please check the path.")
        return

    # Input for bank names
    bank_names_input = st.text_area("Enter Bank Names (comma-separated)", "")
    if bank_names_input:
        bank_names = [name.strip() for name in bank_names_input.split(',')]

        # Generate the document
        if st.button("Generate Word File"):
            bank_details = fetch_emails_from_excel(bank_names)
            if bank_details:
                word_file = create_word_file(bank_details)
                
                # Provide a download button
                st.download_button(
                    label="Download Word File",
                    data=word_file,
                    file_name="Banks_Email_Nodal.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
